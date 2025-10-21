"""
Word Document Generator - Core Engine

This module provides comprehensive Word document generation capabilities
using python-docx library with precise formatting control.

Author: ZTL数智化作战中心
Version: 1.0.0
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
from typing import Dict, List, Any, Optional, Union
from datetime import datetime
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class WordDocumentBuilder:
    """
    Advanced Word document builder with fluent API for precise formatting control.

    Example:
        >>> builder = WordDocumentBuilder()
        >>> builder.add_title("My Document")
        >>> builder.add_heading("Chapter 1", level=1)
        >>> builder.add_paragraph("Content here...")
        >>> builder.save("output.docx")
    """

    def __init__(self, template_path: Optional[str] = None):
        """
        Initialize document builder.

        Args:
            template_path: Optional path to .docx template file
        """
        if template_path and Path(template_path).exists():
            self.doc = Document(template_path)
            logger.info(f"Loaded template: {template_path}")
        else:
            self.doc = Document()
            logger.info("Created new blank document")

        self.current_section = self.doc.sections[0]

    def add_title(self, text: str, align: str = "left") -> "WordDocumentBuilder":
        """
        Add document title with Title style.

        Args:
            text: Title text
            align: Alignment (left, center, right)

        Returns:
            Self for method chaining
        """
        title = self.doc.add_heading(text, level=0)
        title.alignment = self._get_alignment(align)
        logger.info(f"Added title: {text}")
        return self

    def add_heading(self, text: str, level: int = 1, align: str = "left") -> "WordDocumentBuilder":
        """
        Add heading with specified level.

        Args:
            text: Heading text
            level: Heading level (1-6)
            align: Alignment (left, center, right)

        Returns:
            Self for method chaining
        """
        if not 1 <= level <= 6:
            raise ValueError(f"Heading level must be 1-6, got {level}")

        heading = self.doc.add_heading(text, level=level)
        heading.alignment = self._get_alignment(align)
        logger.info(f"Added heading (level {level}): {text}")
        return self

    def add_paragraph(
        self,
        text: str,
        style: Optional[str] = None,
        align: str = "left",
        font_name: Optional[str] = None,
        font_size: Optional[int] = None,
        bold: bool = False,
        italic: bool = False,
        color: Optional[str] = None,
        first_line_indent: Optional[float] = None
    ) -> "WordDocumentBuilder":
        """
        Add paragraph with rich formatting options.

        Args:
            text: Paragraph text
            style: Paragraph style name
            align: Alignment (left, center, right, justify)
            font_name: Font family name
            font_size: Font size in points
            bold: Bold text
            italic: Italic text
            color: RGB color in hex format (e.g., "FF0000" for red)
            first_line_indent: First line indent in inches

        Returns:
            Self for method chaining
        """
        para = self.doc.add_paragraph(text, style=style)
        para.alignment = self._get_alignment(align)

        # Apply formatting to all runs
        for run in para.runs:
            if font_name:
                run.font.name = font_name
            if font_size:
                run.font.size = Pt(font_size)
            if bold:
                run.font.bold = True
            if italic:
                run.font.italic = True
            if color:
                run.font.color.rgb = RGBColor.from_string(color)

        # Set first line indent
        if first_line_indent:
            para.paragraph_format.first_line_indent = Inches(first_line_indent)

        logger.info(f"Added paragraph: {text[:50]}...")
        return self

    def add_bullet_list(self, items: List[str], style: str = "List Bullet") -> "WordDocumentBuilder":
        """
        Add bullet point list.

        Args:
            items: List of bullet point texts
            style: List style name

        Returns:
            Self for method chaining
        """
        for item in items:
            self.doc.add_paragraph(item, style=style)
        logger.info(f"Added bullet list with {len(items)} items")
        return self

    def add_numbered_list(self, items: List[str], style: str = "List Number") -> "WordDocumentBuilder":
        """
        Add numbered list.

        Args:
            items: List of numbered item texts
            style: List style name

        Returns:
            Self for method chaining
        """
        for item in items:
            self.doc.add_paragraph(item, style=style)
        logger.info(f"Added numbered list with {len(items)} items")
        return self

    def add_table(
        self,
        headers: List[str],
        rows: List[List[str]],
        style: str = "Light Grid Accent 1"
    ) -> "WordDocumentBuilder":
        """
        Add table with headers and data rows.

        Args:
            headers: Column header texts
            rows: List of row data (each row is a list of cell values)
            style: Table style name

        Returns:
            Self for method chaining
        """
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = style

        # Add headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Make header bold
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Add data rows
        for i, row_data in enumerate(rows, start=1):
            row_cells = table.rows[i].cells
            for j, cell_value in enumerate(row_data):
                row_cells[j].text = str(cell_value)

        logger.info(f"Added table: {len(headers)} columns × {len(rows)} rows")
        return self

    def add_image(
        self,
        image_path: str,
        width_inches: Optional[float] = None,
        height_inches: Optional[float] = None,
        caption: Optional[str] = None
    ) -> "WordDocumentBuilder":
        """
        Add image to document.

        Args:
            image_path: Path to image file
            width_inches: Image width in inches
            height_inches: Image height in inches
            caption: Optional image caption

        Returns:
            Self for method chaining
        """
        if not Path(image_path).exists():
            logger.error(f"Image not found: {image_path}")
            raise FileNotFoundError(f"Image not found: {image_path}")

        # Add image
        if width_inches and height_inches:
            self.doc.add_picture(image_path, width=Inches(width_inches), height=Inches(height_inches))
        elif width_inches:
            self.doc.add_picture(image_path, width=Inches(width_inches))
        elif height_inches:
            self.doc.add_picture(image_path, height=Inches(height_inches))
        else:
            self.doc.add_picture(image_path)

        # Add caption if provided
        if caption:
            self.add_paragraph(caption, align="center", italic=True, font_size=10)

        logger.info(f"Added image: {image_path}")
        return self

    def add_page_break(self) -> "WordDocumentBuilder":
        """Add page break."""
        self.doc.add_page_break()
        logger.info("Added page break")
        return self

    def set_page_margins(
        self,
        top: float = 1.0,
        bottom: float = 1.0,
        left: float = 1.0,
        right: float = 1.0
    ) -> "WordDocumentBuilder":
        """
        Set page margins.

        Args:
            top: Top margin in inches
            bottom: Bottom margin in inches
            left: Left margin in inches
            right: Right margin in inches

        Returns:
            Self for method chaining
        """
        self.current_section.top_margin = Inches(top)
        self.current_section.bottom_margin = Inches(bottom)
        self.current_section.left_margin = Inches(left)
        self.current_section.right_margin = Inches(right)
        logger.info(f"Set margins: T={top} B={bottom} L={left} R={right}")
        return self

    def add_header(self, text: str) -> "WordDocumentBuilder":
        """
        Add text to document header.

        Args:
            text: Header text

        Returns:
            Self for method chaining
        """
        header = self.current_section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.text = text
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logger.info(f"Added header: {text}")
        return self

    def add_footer(self, text: str) -> "WordDocumentBuilder":
        """
        Add text to document footer.

        Args:
            text: Footer text (use {page} for page number, {total_pages} for total)

        Returns:
            Self for method chaining
        """
        footer = self.current_section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.text = text
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logger.info(f"Added footer: {text}")
        return self

    def save(self, output_path: str) -> Dict[str, Any]:
        """
        Save document to file.

        Args:
            output_path: Output file path

        Returns:
            Result dictionary with success status and file path
        """
        try:
            # Ensure output directory exists
            output_file = Path(output_path)
            output_file.parent.mkdir(parents=True, exist_ok=True)

            # Save document
            self.doc.save(str(output_file))

            logger.info(f"Document saved: {output_path}")
            return {
                "success": True,
                "file_path": str(output_file.absolute()),
                "size_bytes": output_file.stat().st_size
            }
        except Exception as e:
            logger.error(f"Failed to save document: {e}")
            return {
                "success": False,
                "error": str(e)
            }

    def _get_alignment(self, align: str) -> int:
        """Convert alignment string to WD_ALIGN_PARAGRAPH constant."""
        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        return alignment_map.get(align.lower(), WD_ALIGN_PARAGRAPH.LEFT)


def create_document(
    title: str,
    content: List[Dict[str, Any]],
    output_path: str,
    author: Optional[str] = None,
    template_path: Optional[str] = None
) -> Dict[str, Any]:
    """
    Create Word document from structured content.

    Args:
        title: Document title
        content: List of content elements (see examples below)
        output_path: Output file path
        author: Document author (optional)
        template_path: Template file path (optional)

    Content element examples:
        {"type": "heading", "level": 1, "text": "Chapter 1"}
        {"type": "paragraph", "text": "Content...", "bold": True}
        {"type": "bullet_list", "items": ["Item 1", "Item 2"]}
        {"type": "numbered_list", "items": ["Step 1", "Step 2"]}
        {"type": "table", "headers": [...], "rows": [[...]]}
        {"type": "image", "path": "image.png", "width": 6}
        {"type": "page_break"}

    Returns:
        Result dictionary with success status and file path
    """
    try:
        builder = WordDocumentBuilder(template_path)

        # Add title
        builder.add_title(title, align="center")

        # Add author and date if provided
        if author:
            builder.add_paragraph(f"作者: {author}", align="right", font_size=10)
        builder.add_paragraph(f"日期: {datetime.now().strftime('%Y-%m-%d')}", align="right", font_size=10)
        builder.add_paragraph("")  # Empty line

        # Process content elements
        for element in content:
            element_type = element.get("type")

            if element_type == "heading":
                builder.add_heading(
                    element["text"],
                    level=element.get("level", 1),
                    align=element.get("align", "left")
                )

            elif element_type == "paragraph":
                builder.add_paragraph(
                    element["text"],
                    style=element.get("style"),
                    align=element.get("align", "left"),
                    font_name=element.get("font_name"),
                    font_size=element.get("font_size"),
                    bold=element.get("bold", False),
                    italic=element.get("italic", False),
                    color=element.get("color")
                )

            elif element_type == "bullet_list":
                builder.add_bullet_list(element["items"])

            elif element_type == "numbered_list":
                builder.add_numbered_list(element["items"])

            elif element_type == "table":
                builder.add_table(
                    element["headers"],
                    element["rows"],
                    style=element.get("style", "Light Grid Accent 1")
                )

            elif element_type == "image":
                builder.add_image(
                    element["path"],
                    width_inches=element.get("width"),
                    height_inches=element.get("height"),
                    caption=element.get("caption")
                )

            elif element_type == "page_break":
                builder.add_page_break()

        # Save document
        return builder.save(output_path)

    except Exception as e:
        logger.error(f"Failed to create document: {e}")
        return {
            "success": False,
            "error": str(e)
        }


def create_from_template(
    template_type: str,
    data: Dict[str, Any],
    output_path: str
) -> Dict[str, Any]:
    """
    Create document from predefined template.

    Args:
        template_type: Template type (report, proposal, contract, meeting-minutes, manual, letter)
        data: Template data dictionary
        output_path: Output file path

    Returns:
        Result dictionary with success status and file path
    """
    try:
        # Use absolute import instead of relative
        import templates

        template = templates.TemplateRegistry.get_template(template_type)
        if not template:
            raise ValueError(f"Unknown template type: {template_type}")

        content = template.render(data)
        return create_document(
            title=data.get("title", "Untitled"),
            content=content,
            output_path=output_path,
            author=data.get("author")
        )

    except Exception as e:
        logger.error(f"Failed to create from template: {e}")
        return {
            "success": False,
            "error": str(e)
        }


# Convenience function for quick document creation
def quick_create(
    title: str,
    sections: List[tuple],
    output_path: str
) -> Dict[str, Any]:
    """
    Quick document creation with simple section tuples.

    Args:
        title: Document title
        sections: List of (heading, content) tuples
        output_path: Output file path

    Example:
        >>> quick_create(
        ...     "My Report",
        ...     [
        ...         ("Introduction", "This is the intro..."),
        ...         ("Analysis", "Here are the findings..."),
        ...         ("Conclusion", "In summary...")
        ...     ],
        ...     "output/report.docx"
        ... )

    Returns:
        Result dictionary with success status and file path
    """
    content = []
    for heading, text in sections:
        content.append({"type": "heading", "level": 1, "text": heading})
        content.append({"type": "paragraph", "text": text})

    return create_document(title, content, output_path)
